from pathlib import Path
from datetime import date
import uuid
import textwrap
from typing import Dict, Any, List

from PySide6.QtCore import Qt, QUrl, QMimeData
from PySide6.QtGui import QDesktopServices, QGuiApplication
from PySide6.QtWidgets import (
    QDialog, QFormLayout, QTextEdit, QLineEdit, QPushButton,
    QVBoxLayout, QWidget, QHBoxLayout, QMessageBox, QLabel
)

from docx import Document as Docx
from docx.shared import Pt

from employ_toolkit.core.models import Document
from employ_toolkit.core.storage import get_session

OUTPUT_DIR = Path("workspace")
OUTPUT_DIR.mkdir(exist_ok=True)

CHATGPT_URL = "https://chatgpt.com/g/g-cStsvQbjd-cv-writer-the-cv-expert"
PROMPT_TXT = textwrap.dedent("""\
Actúa como experto en redacción de CV. Analiza mi currículum y devuélvelo
optimizado para ATS (palabras clave, logros cuantificados, perfil de impacto).
*** Pega abajo tu CV completo ***""")

# --------------------------------------------------------------------------- #
class CVForm(QDialog):
    """Formulario para construir un CV optimizado (genera solo DOCX)."""

    def __init__(self, client, parent=None):
        super().__init__(parent)
        self.client = client
        self.setWindowTitle(f"CV Optimizado – {client.full_name}")
        self.resize(620, 720)

        main = QVBoxLayout(self)

        # ---------- Botón GPT ----------
        gpt_btn = QPushButton("Abrir ChatGPT · CV Writer")
        gpt_btn.clicked.connect(self._open_gpt)
        main.addWidget(gpt_btn)

        # ---------- Formulario ----------
        form = QFormLayout()

        self.name      = QLineEdit(self.client.full_name)
        self.contact   = QLineEdit()  # teléfono · email · ciudad
        self.title     = QLineEdit()
        self.summary   = QTextEdit(); self.summary.setFixedHeight(60)
        self.skills    = QTextEdit(); self.skills.setFixedHeight(40)
        self.exp       = QTextEdit(); self.exp.setFixedHeight(120)
        self.edu       = QTextEdit(); self.edu.setFixedHeight(60)
        self.certs     = QTextEdit(); self.certs.setFixedHeight(40)
        self.langs     = QLineEdit()
        self.extras    = QTextEdit(); self.extras.setFixedHeight(40)

        form.addRow("Nombre completo", self.name)
        form.addRow("Datos contacto (tel/e-mail)", self.contact)
        form.addRow("Título profesional", self.title)
        form.addRow("Resumen / Perfil", self.summary)
        form.addRow("Competencias clave (Hard / Soft Skills)", self.skills)
        form.addRow(QLabel("Experiencia (1 línea por puesto):"), self.exp)
        form.addRow("Educación", self.edu)
        form.addRow("Certificaciones", self.certs)
        form.addRow("Idiomas", self.langs)
        form.addRow("Premios / Publicaciones", self.extras)

        main.addLayout(form)

        gen_btn = QPushButton("Generar DOCX")
        gen_btn.clicked.connect(self._generate)
        main.addWidget(gen_btn, alignment=Qt.AlignRight)

    # ------------ Helpers -------------
    def _open_gpt(self):
        QDesktopServices.openUrl(QUrl(CHATGPT_URL))
        md = QMimeData(); md.setText(PROMPT_TXT)
        QGuiApplication.clipboard().setMimeData(md)
        QMessageBox.information(self, "Prompt copiado",
                                "Se abrió ChatGPT y el prompt está en tu portapapeles.")

    def _collect(self) -> Dict[str, str]:
        return dict(
            nombre=self.name.text().strip(),
            contacto=self.contact.text().strip(),
            titulo=self.title.text().strip(),
            resumen=self.summary.toPlainText().strip(),
            skills=self.skills.toPlainText().strip(),
            experiencia=self.exp.toPlainText().strip(),
            educacion=self.edu.toPlainText().strip(),
            certs=self.certs.toPlainText().strip(),
            idiomas=self.langs.text().strip(),
            extras=self.extras.toPlainText().strip(),
        )

    # ------------ Generar -------------
    def _generate(self):
        data = self._collect()
        if not data["titulo"] or not data["resumen"]:
            QMessageBox.warning(self, "Faltan datos",
                                "Título y Resumen son obligatorios.")
            return

        docx_path = self._build_docx(data)
        self._save_in_db(docx_path)

        QMessageBox.information(self, "CV creado", f"DOCX: {docx_path.name}")
        self.accept()

    # ------------ DOCX builder --------
    def _build_docx(self, d: Dict[str, str]) -> Path:
        uid = uuid.uuid4().hex[:6]
        path = OUTPUT_DIR / f"{self.client.full_name}_{uid}_cv.docx"

        doc = Docx()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Encabezado nombre + contacto
        doc.add_heading(d["nombre"], level=0)
        if d["contacto"]:
            p = doc.add_paragraph(d["contacto"])
            p.runs[0].italic = True

        doc.add_heading(d["titulo"], level=1)
        doc.add_paragraph(d["resumen"])

        def h1(txt): doc.add_heading(txt, level=1)

        h1("Competencias clave")
        for sk in d["skills"].split(","):
            if sk.strip():
                doc.add_paragraph(sk.strip(), style="List Bullet")

        h1("Experiencia")
        for line in d["experiencia"].splitlines():
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")

        h1("Educación")
        for line in d["educacion"].splitlines():
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")

        if d["certs"]:
            h1("Certificaciones")
            for line in d["certs"].splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")

        if d["idiomas"]:
            h1("Idiomas")
            doc.add_paragraph(d["idiomas"])

        if d["extras"]:
            h1("Premios / Publicaciones")
            for line in d["extras"].splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")

        doc.save(path)
        return path

    # ------------ DB register ----------
    def _save_in_db(self, docx_path: Path):
        with get_session() as s:
            s.add(Document(
                client_id=self.client.id,
                module=2, doc_type="cv_docx",
                path=str(docx_path), created_at=date.today()
            ))
            s.commit()
