"""
cv_builder.py
-------------
• Analiza CV actual (PDF/DOCX) y role target
• Llama a ChatGPT (mock en dev) para entregar JSON con campos normalizados
• Renderiza PDF y DOCX finales
"""

from pathlib import Path
import uuid, json
from datetime import date

OUTPUT_DIR = Path("workspace")
OUTPUT_DIR.mkdir(exist_ok=True)

# ---------- 1. Mock del análisis GPT ----------
def analyse_cv_with_gpt(original_text: str, role: str) -> dict:
    """
    Mock: devuelve un diccionario con campos estándar que rellenan el formulario.
    En producción esta función llamará a openai.ChatCompletion.
    """
    return {
        "role_target": role,
        "summary": f"Profesional orientado a {role} con 5+ años...",
        "hard_skills": ["Python", "SQL", "Power BI"],
        "soft_skills": ["Comunicación", "Liderazgo"],
        "achievements": [
            {"text": "Incrementé las ventas 25 %", "metric": "25 %"}
        ],
        "experience": [
            {
                "company": "ACME",
                "title": role,
                "dates": "2021-2024",
                "bullets": ["Desarrollé dashboards…", "Automaticé ETL…"]
            }
        ],
        "education": [{"degree": "Ingeniero", "year": "2018"}],
        "keywords_missing": ["ETL", "Docker"],
    }

# ---------- 2. Renderizadores ----------
def _build_docx(data: dict, filename: Path):
    from docx import Document
    doc = Document()
    doc.add_heading(data["role_target"], level=1)
    doc.add_paragraph(data["summary"])
    doc.save(filename)

def _build_pdf(data: dict, filename: Path):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    doc = SimpleDocTemplate(str(filename), pagesize=LETTER)
    doc.build([Paragraph(data["summary"], getSampleStyleSheet()["Normal"])])

# ---------- 3. Función principal ----------
def generate_cv_files(client, data: dict) -> dict[str, Path]:
    """
    data := formulario completo (campos mostrados en GUI)
    Devuelve {"pdf": Path, "docx": Path}
    """
    stem = f"{client.full_name}_{uuid.uuid4().hex[:6]}_cv"
    pdf = OUTPUT_DIR / f"{stem}.pdf"
    docx = OUTPUT_DIR / f"{stem}.docx"

    _build_pdf(data, pdf)
    _build_docx(data, docx)

    return {"pdf": pdf, "docx": docx}
