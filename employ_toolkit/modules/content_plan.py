# employ_toolkit/modules/content_plan.py
import uuid
from pathlib import Path
from datetime import date, timedelta

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

OUTPUT_DIR = Path("workspace")
OUTPUT_DIR.mkdir(exist_ok=True)


def _week_start(reference: date, n: int) -> date:
    """Devuelve el lunes de la semana n (0-based) a partir de 'reference'."""
    ref_monday = reference + timedelta(days=-reference.weekday())  # lunes actual
    return ref_monday + timedelta(weeks=n)


def generate_content_plan(client, params: dict) -> dict:
    """
    Genera un DOCX resumen + XLSX calendario.
    params = {pilares: [...], freq: int, formatos:[...], semanas:int}
    Devuelve {'docx': Path, 'xlsx': Path}
    """
    # ---------- DOCX RESUMEN ----------
    docx_path = OUTPUT_DIR / f"{client.full_name}_{uuid.uuid4().hex[:6]}_plan.docx"
    doc = Document()
    doc.add_heading(f"Parrilla de Contenidos – {client.full_name}", 0)
    doc.add_paragraph(f"Pilares: {', '.join(params['pilares'])}")
    doc.add_paragraph(f"Frecuencia: {params['freq']} publicaciones / semana")
    doc.add_paragraph(f"Formatos permitidos: {', '.join(params['formatos'])}")
    doc.add_paragraph(f"Duración piloto: {params['semanas']} semanas")
    doc.save(docx_path)

    # ---------- XLSX CALENDARIO ----------
    xlsx_path = OUTPUT_DIR / f"{client.full_name}_{uuid.uuid4().hex[:6]}_plan.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Calendario"

    header = ["Fecha", "Pilar", "Formato", "Hook / Título", "Idea / Copy (CTA)"]
    ws.append(header)
    for col in "ABCDE":
        ws[f"{col}1"].font = Font(bold=True)
        ws[f"{col}1"].alignment = Alignment(horizontal="center")

    today = date.today()
    pillars = params["pilares"]
    formats = params["formatos"]
    freq = params["freq"]
    semanas = params["semanas"]

    # Ejemplos para la primera semana
    hooks = ["¿Sabías que…?", "Tip rápido:", "Historia personal:", "Dato revelador:", "Pregunta al público:"]
    ctas  = ["Comenta tu experiencia.", "Guarda este post.", "Compártelo con tu red."]

    row_idx = 2
    for w in range(semanas):
        monday = _week_start(today, w)
        for i in range(freq):
            fecha = monday + timedelta(days=i % 5)  # de lunes a viernes
            pilar = pillars[i % len(pillars)]
            formato = formats[i % len(formats)]

            if w == 0:  # primera semana con ejemplo
                hook = hooks[i % len(hooks)]
                cta  = ctas[i % len(ctas)]
                titulo = f"{hook} ({pilar})"
                copy   = f"Desarrollo del tema sobre {pilar.lower()}. {cta}"
            else:
                titulo = ""
                copy   = ""

            ws.append([str(fecha), pilar, formato, titulo, copy])
            row_idx += 1

    # Ajuste de columnas
    for col, width in zip("ABCDE", [12, 16, 14, 40, 60]):
        ws.column_dimensions[col].width = width

    wb.save(xlsx_path)
    return {"docx": docx_path, "xlsx": xlsx_path}

